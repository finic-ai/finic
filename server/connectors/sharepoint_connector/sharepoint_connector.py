import requests
import os
import json
from typing import Dict, List, Optional
from models.models import (
    AppConfig,
    Document,
    ConnectorId,
    DocumentConnector,
    AuthorizationResult,
    ConnectionFilter,
)
from appstatestore.statestore import StateStore
import base64
from models.api import GetDocumentsResponse
from urllib.parse import urlencode, urlunparse, urlparse, ParseResult
import tempfile
from docx import Document as DocxDocument
import PyPDF2


def read_docx(file_path):
    doc = DocxDocument(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return "\n".join(full_text)


def read_pdf(file_path):
    with open(file_path, "rb") as file:
        pdf_reader = PyPDF2.PdfFileReader(file)
        count = pdf_reader.numPages
        content = ""
        for i in range(count):
            page = pdf_reader.getPage(i)
            content += page.extract_text()
    return content


def get_items(drive_id, item_id, headers):
    response = requests.get(
        f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item_id}/children",
        headers=headers,
    )
    response_data = response.json()
    return response_data.get("value")


def get_content(doc_title, doc_url, headers):
    response = requests.get(doc_url, headers=headers)
    with tempfile.NamedTemporaryFile(suffix=".tmp", delete=False) as temp:
        temp.write(response.content)
        temp_file_path = temp.name

    doc_content = ""
    if doc_title.endswith(".docx"):
        doc_content = read_docx(temp_file_path)
    elif doc_title.endswith(".pdf"):
        doc_content = read_pdf(temp_file_path)
    os.remove(temp_file_path)

    return doc_content


def parse_items(items, drive_id, headers, account_id):
    docs = []
    for item in items:
        if item.get("folder", None):  # if the item is a folder
            # recursively get the documents in the folder
            sub_items = get_items(drive_id, item["id"], headers)
            docs.extend(parse_items(sub_items, drive_id, headers, account_id))
        else:
            # if the item is a file
            doc_title = item["name"]
            doc_url = item["@microsoft.graph.downloadUrl"]
            doc_content = get_content(doc_title, doc_url, headers)
            web_url = item["webUrl"]
            doc = Document(
                title=doc_title,
                uri=web_url,
                content=doc_content,
                connector_id=ConnectorId.sharepoint,
                account_id=account_id,
            )
            docs.append(doc)
    return docs


class SharepointConnector(DocumentConnector):
    connector_id: ConnectorId = ConnectorId.sharepoint
    config: AppConfig
    headers: Dict = {}

    def __init__(self, config: AppConfig):
        super().__init__(config=config)

    async def authorize_api_key(self) -> AuthorizationResult:
        pass

    async def authorize(
        self, account_id: str, auth_code: Optional[str], metadata: Optional[Dict]
    ) -> AuthorizationResult:
        connector_credentials = StateStore().get_connector_credential(
            self.connector_id, self.config
        )
        try:
            client_id = connector_credentials["client_id"]
            client_secret = connector_credentials["client_secret"]
            redirect_uri = "https://link.psychic.dev/oauth/redirect"
        except Exception as e:
            raise Exception("Connector is not enabled")

        if not auth_code:
            auth_url = (
                f"https://login.microsoftonline.com/organizations/oauth2/v2.0/authorize"
            )
            params = {
                "response_type": "code",
                "redirect_uri": redirect_uri,
                "client_id": client_id,
                "scope": "offline_access https://graph.microsoft.com/.default",
                "prompt": "consent",
            }
            encoded_params = urlencode(params)
            url_parts = urlparse(auth_url)
            auth_url = ParseResult(
                scheme=url_parts.scheme,
                netloc=url_parts.netloc,
                path=url_parts.path,
                params=url_parts.params,
                query=encoded_params,
                fragment=url_parts.fragment,
            ).geturl()
            return AuthorizationResult(authorized=False, auth_url=auth_url)

        try:
            # encode in base 64
            headers = {
                "Content-Type": "application/x-www-form-urlencoded",
            }

            data = {
                "client_id": client_id,
                "scope": "https://graph.microsoft.com/.default",
                "code": auth_code,
                "redirect_uri": redirect_uri,
                "grant_type": "authorization_code",
                "client_secret": client_secret,
            }

            response = requests.post(
                f"https://login.microsoftonline.com/common/oauth2/v2.0/token",
                headers=headers,
                data=data,
            )

            creds = response.json()

            creds_string = json.dumps(creds)

            headers = {"Authorization": "Bearer " + creds.get("access_token")}

            response = requests.get(
                "https://graph.microsoft.com/v1.0/organization", headers=headers
            )
            response_json = response.json()

            org_name = None

            orgs = response_json.get("value")

            if len(orgs) > 0:
                org_name = orgs[0].get("displayName")

        except Exception as e:
            print(e)
            raise Exception("Unable to get access token with code")

        new_connection = StateStore().add_connection(
            config=self.config,
            credential=creds_string,
            connector_id=self.connector_id,
            account_id=account_id,
            metadata={"org": org_name},
        )
        return AuthorizationResult(authorized=True, connection=new_connection)

    async def get_sections(self) -> List[str]:
        pass

    async def load(self, connection_filter: ConnectionFilter) -> GetDocumentsResponse:
        connector_credentials = StateStore().get_connector_credential(
            self.connector_id, self.config
        )
        try:
            client_id = connector_credentials["client_id"]
            client_secret = connector_credentials["client_secret"]
            redirect_uri = "https://link.psychic.dev/oauth/redirect"
        except Exception as e:
            raise Exception("Connector is not enabled")

        account_id = connection_filter.account_id
        uris = connection_filter.uris
        section_filter = connection_filter.section_filter_id
        connection = StateStore().load_credentials(
            self.config, self.connector_id, account_id
        )
        credential_string = connection.credential
        credential_json = json.loads(credential_string)

        print(credential_json)

        access_token = credential_json.get("access_token")

        # check if access token is expired and refresh it

        headers = {"Authorization": "Bearer " + access_token}

        response = requests.get(
            "https://graph.microsoft.com/v1.0/sites?search=*", headers=headers
        )
        if response.status_code == 401:
            # refresh token
            refresh_token = credential_json.get("refresh_token")
            data = {
                "client_id": client_id,
                "scope": "https://graph.microsoft.com/.default",
                "refresh_token": refresh_token,
                "grant_type": "refresh_token",
                "client_secret": client_secret,
            }
            response = requests.post(
                f"https://login.microsoftonline.com/common/oauth2/v2.0/token",
                headers=headers,
                data=data,
            )
            creds = response.json()
            creds_string = json.dumps(creds)
            new_connection = StateStore().add_connection(
                config=self.config,
                credential=creds_string,
                connector_id=self.connector_id,
                account_id=account_id,
                metadata={},
            )
            access_token = creds.get("access_token")
            headers = {"Authorization": "Bearer " + access_token}
            response = requests.get(
                "https://graph.microsoft.com/v1.0/sites?search=*", headers=headers
            )
        response_data = response.json()
        all_sites = response_data.get("value")

        all_drives = []
        if all_sites:
            for site in all_sites:
                site_id = site.get("id")
                response = requests.get(
                    f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives",
                    headers=headers,
                )
                response_data = response.json()
                drives = response_data.get("value")
                all_drives.extend(drives)

        all_docs = []
        if all_drives:
            for drive in all_drives:
                drive_id = drive.get("id")
                items = get_items(
                    drive_id, "root", headers
                )  # getting the children of the root
                all_docs.extend(parse_items(items, drive_id, headers, account_id))

        return GetDocumentsResponse(documents=all_docs, next_page_cursor=None)
