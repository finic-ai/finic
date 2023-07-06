import requests
import json
from typing import Dict, List, Optional
from models.models import (
    AppConfig,
    Document,
    ConnectorId,
    DataConnector,
    AuthorizationResult,
    ConnectionFilter,
)
from models.api import GetDocumentsRequest, GetDocumentsResponse
from urllib.parse import urlencode, urlunparse, urlparse, ParseResult
from appstatestore.statestore import StateStore
from io import BytesIO
from PyPDF2 import PdfReader
import docx
from enum import Enum


BASE_URL = "https://api.dropboxapi.com"


class DropboxError(str, Enum):
    expired_access_token = "expired_access_token"
    unknown_error = "unknown_error"


class DropboxConnector(DataConnector):
    connector_id: ConnectorId = ConnectorId.dropbox
    config: AppConfig

    def __init__(self, config: AppConfig):
        super().__init__(config=config)

    def check_valid_access_token(self, access_token: str) -> bool:
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json",
        }

        json_data = {
            "query": "foo",
        }
        response = requests.post(
            f"{BASE_URL}/2/check/user", headers=headers, json=json_data
        )
        if response.status_code == 200:
            return True
        else:
            return False

    async def authorize_api_key(self) -> AuthorizationResult:
        pass

    async def authorize(
        self, account_id: str, auth_code: Optional[str], metadata: Dict
    ) -> AuthorizationResult:
        connector_credentials = StateStore().get_connector_credential(
            self.connector_id, self.config
        )
        try:
            client_id = connector_credentials["client_id"]
            client_secret = connector_credentials["client_secret"]
            redirect_uri = "https://link.psychic.dev/oauth/redirect"
        except Exception as e:
            raise Exception("Connector is not enabled")

        if not auth_code:
            auth_url = "https://www.dropbox.com/oauth2/authorize"
            scopes = "account_info.read files.metadata.read files.content.read"
            params = {
                "response_type": "code",
                "redirect_uri": redirect_uri,
                "client_id": client_id,
                "scope": scopes,
                "token_access_type": "offline",
            }
            encoded_params = urlencode(params)
            url_parts = urlparse(auth_url)
            auth_url = ParseResult(
                scheme=url_parts.scheme,
                netloc=url_parts.netloc,
                path=url_parts.path,
                params=url_parts.params,
                query=encoded_params,
                fragment=url_parts.fragment,
            ).geturl()
            return AuthorizationResult(authorized=False, auth_url=auth_url)

        try:
            data = {
                "code": auth_code,
                "grant_type": "authorization_code",
                "client_id": client_id,
                "client_secret": client_secret,
                "redirect_uri": redirect_uri,
            }

            response = requests.post("https://api.dropbox.com/oauth2/token", data=data)
            response_data = response.json()

            creds_string = json.dumps(response_data)
            metadata = {
                "dropbox_account_id": response_data["account_id"],
            }
            new_connection = StateStore().add_connection(
                config=self.config,
                credential=creds_string,
                connector_id=self.connector_id,
                account_id=account_id,
                metadata=metadata,
            )
            return AuthorizationResult(authorized=True, connection=new_connection)
        except Exception as e:
            print(e)
            raise Exception("Unable to get access token with code")

    def get_new_access_token(self, refresh_token: str):
        connector_credentials = StateStore().get_connector_credential(
            self.connector_id, self.config
        )
        client_id = connector_credentials["client_id"]
        client_secret = connector_credentials["client_secret"]
        data = {
            "grant_type": "refresh_token",
            "refresh_token": refresh_token,
            "client_id": client_id,
            "client_secret": client_secret,
        }

        response = requests.post("https://api.dropbox.com/oauth2/token", data=data)
        if response.status_code == 200:
            res_data = response.json()
            return res_data.get("access_token")
        else:
            return None

    def get_all_files_under_folder(self, access_token: str):
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Content-Type": "application/json",
        }

        json_data = {
            "include_deleted": False,
            "include_has_explicit_shared_members": False,
            "include_media_info": False,
            "include_mounted_folders": True,
            "include_non_downloadable_files": True,
            "path": f"",
            "recursive": True,
        }

        res = requests.post(
            f"{BASE_URL}/2/files/list_folder", headers=headers, json=json_data
        )
        data = res.json()

        if "error" in data:
            error = data.get("error")
            if error.get(".tag") == DropboxError.expired_access_token:
                return {"error": DropboxError.expired_access_token}
            else:
                return {"error": DropboxError.unknown_error}

        files = data.get("entries")

        file_metadata = []
        for doc in files:
            if doc.get(".tag") == "file":
                document_name = doc.get("name")
                file_type = document_name.split(".")[-1]
                if doc.get("is_downloadable"):
                    document_path = doc.get("path_lower")
                    file_metadata.append((document_name, document_path, file_type))

        return {"results": file_metadata}

    def extract_text_from_document(
        self, access_token: str, document_path: str, file_type: str
    ):
        path = {"path": document_path}
        path = json.dumps(path)
        headers = {
            "Authorization": f"Bearer {access_token}",
            "Dropbox-API-Arg": path,
        }

        response = requests.post(
            "https://content.dropboxapi.com/2/files/download", headers=headers
        )
        if response.status_code == 200:
            content = response.content
            if content:
                text = ""
                if file_type == "pdf":
                    bytes_data = BytesIO(content)
                    reader = PdfReader(bytes_data)
                    for num in range(len(reader.pages)):
                        page = reader.pages[num]
                        page_text = page.extract_text()
                        text += page_text

                elif file_type == "docx" or file_type == "doc":
                    bytes_data = BytesIO(content)
                    doc = docx.Document(bytes_data)
                    for p in doc.paragraphs:
                        text += p.text

                elif file_type == "txt":
                    text = content.decode("utf-8")
                return text
        return None

    async def get_sections(self) -> List[str]:
        pass

    async def load(self, connection_filter: ConnectionFilter) -> GetDocumentsResponse:
        account_id = connection_filter.account_id
        uris = connection_filter.uris
        section_filter = connection_filter.section_filter_id
        connection = StateStore().load_credentials(
            self.config, self.connector_id, account_id=account_id
        )
        credential_json = json.loads(connection.credential)

        documents: List[Document] = []

        access_token = credential_json.get("access_token")
        refresh_token = credential_json.get("refresh_token")

        try:
            result = self.get_all_files_under_folder(access_token)
            if "error" in result:
                if result.get("error") == DropboxError.expired_access_token:
                    access_token = self.get_new_access_token(refresh_token)
                    result = self.get_all_files_under_folder(access_token)
                else:
                    raise Exception("Error in getting files")
            files = result.get("results")
        except Exception as e:
            raise e

        for file_name, file_path, file_type in files:
            text = self.extract_text_from_document(access_token, file_path, file_type)
            if text:
                documents.append(
                    Document(
                        title=file_name,
                        content=text,
                        connector_id=self.connector_id,
                        account_id=account_id,
                        uri=file_path,
                    )
                )

        return GetDocumentsResponse(documents=documents)
