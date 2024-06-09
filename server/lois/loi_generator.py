import os

from models.models import (
    LOI,
)

from typing import List, Any, Optional, Tuple
from database import Database
import json
import re
import base64
import uuid
import io
import zipfile
from docx import Document
from docx.shared import Inches

test_loi = LOI(
    id="1",
    status="draft",
    created_by="1234",
    business_name="Test Business",
    buyer_name="Test Buyer",
    legal_entity="Test Legal Entity",
    biz_revenue=1000000,
    biz_ebitda=500000,
    financials_period="ttm",
    purchase_price=10000000,
    note_percent=0.1,
    note_interest_rate=0.05,
    note_term=5,
    note_standby=0,
    transaction_type="asset",
    earnout_description="Test Earnout description",
    escrow_percent=0.1,
    closing_date="2023-01-01",
    exclusivity_start_date="2022-01-01",
    exclusivity_end_date="2022-02-01",
    termination_fee_type="mutual",
    termination_fee_amount=10000,
    governing_law="ca",
    expiration_date="2022-12-31",

)

def load_loi_template(template_name: str) -> Document:
    document = Document(f"templates/{template_name}.docx")
    return document


def generate_loi_docx():
    document = Document()
    document.add_heading('Letter of Intent', 0)
    document.add_heading('1. Parties', level=1)
    document.add_paragraph(f'This Letter of Intent ("LOI") is entered into as of the date of signature below by and between {test_loi.buyer_name} ("Buyer") and {test_loi.business_name} ("Seller").')
    document.add_heading('2. Transaction', level=1)
    document.add_paragraph(f'Buyer and Seller intend to enter into a transaction (the "Transaction") pursuant to which Buyer will acquire from Seller the assets of the business of {test_loi.business_name} (the "Business").')
    document.add_heading('3. Purchase Price', level=1)
    document.add_paragraph(f'The purchase price for the Business shall be ${test_loi.purchase_price} (the "Purchase Price").')
    document.add_heading('4. Note', level=1)
    document.add_paragraph(f'Buyer shall pay the Purchase Price in the form of a promissory note (the "Note") in the principal amount of ${test_loi.purchase_price * test_loi.note_percent} (the "Principal Amount").')
    document.add_heading('5. Interest Rate', level=1)
    document.add_paragraph(f'The Note shall bear interest at a rate of {test_loi.note_interest_rate * 100}% per annum.')
    document.add_heading('6. Term', level=1)
    document.add_paragraph(f'The Note shall have a term of {test_loi.note_term} years.')
    document.add_heading('7. Standby', level=1)
    document.add_paragraph(f'The Note shall be standby.')
    document.add_heading('8. Transaction Type', level=1)
    document.add_paragraph(f'The Transaction shall be an {test_loi.transaction_type} transaction.')
    document.add_heading('9. Earnout', level=1)
    document.add_paragraph(f'{test_loi.earnout_description}')
    document.add_heading('10. Escrow', level=1)
    document.add_paragraph(f'Buyer shall deposit ${test_loi.purchase_price * test_loi.escrow_percent} in escrow.')
    document.add_heading('11. Closing Date', level=1)
    document.add_paragraph(f'The closing date for the Transaction shall be {test_loi.closing_date}.')
    document.add_heading('12. Exclusivity', level=1)
    document.add_paragraph(f'The exclusivity period for the Transaction shall be from {test_loi.exclusivity_start_date} to {test_loi.exclusivity_end_date}.')
    document.add_heading('13. Termination Fee', level=1)
    document.add_paragraph(f'The termination fee for the Transaction shall be ${test_loi.termination_fee_amount} and shall be mutual.')
    document.add_heading('14. Governing Law', level=1)
    document.add_paragraph(f'The Transaction shall be governed by the laws of the state of {test_loi.governing_law}.')
    document.add_heading('15. Expiration Date', level=1)
    document.add_paragraph(f'This LOI shall expire on {test_loi.expiration_date}.')
    document.add_heading('16. Signature', level=1)
    document.add_paragraph(f'Buyer: ___________________________')
    document.add_paragraph(f'Seller: ___________________________')
    document.save('loi.docx')
    
load_loi_template("loi")