import os
from pydantic import BaseModel
from abc import ABC, abstractmethod
from typing import List, Optional, Dict, Any, Tuple
from datetime import datetime
from enum import Enum
from docx import Document
from docx.document import Document as DocxDocument
from docx.shared import Inches
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER, WD_ALIGN_PARAGRAPH
from models.models import LOI
from lois.loi import LOIDocument
import inflect
import pdb

inflect_engine = inflect.engine()

class TestLOI(LOI):
    business_entity_type: Optional[str] = None
    business_address: Optional[str] = None
    business_state: Optional[str] = None
    buyer_title: Optional[str] = None
    seller_name: Optional[str] = None
    equity_rollover_percent: Optional[float] = None
    bullet_payment_anniversary: Optional[int] = None
    escrow_cap: Optional[int] = None
    escrow_tipping_basket: Optional[int] = None
    due_diligence_items: Optional[List[str]] = None
    exclusivity_days: Optional[int] = None

class StanfordBasicLOI(LOIDocument):

    def __init__(self, loi: TestLOI):
        self.loi = loi

    def construct_docx(self) -> None:
        loi = self.loi
        self.document = Document()

        ### Disclaimer
        p = self.document.add_paragraph()
        p.add_run("Dealwise Advisors LLC").bold = True
        self.document.add_paragraph("This document is based on the original form legal documents provided by Goodwin Procter LLP for use in the Stanford GSB Search Fund Primer, which can be accessed at https://www.gsb.stanford.edu/experience/about/centers-institutes/ces/research/search-funds/primer. Dealwise Advisors LLC has modified the original template for illustrative purposes only, and it should not be used without consulting an attorney. Dealwise Advisors LLC has no affiliation with Goodwin Procter LLP or Stanford Graduate School of Business.")
        p = self.document.add_paragraph()
        p.add_run("Disclaimer of Warranties").bold = True
        self.document.add_paragraph("This document is being provided on an “as is” and “as available” basis. None of Dealwise Advisors LLC, its affiliates, subsidiaries or its or their respective officers, directors, employees or agents (collectively, “Dealwise”) guarantees the accuracy, completeness, timeliness, reliability, suitability or usefulness of any portion of this document. None of the Dealwise Parties warrant that this document will be error-free. None of the Dealwise Parties makes any, and each Dealwise Party hereby disclaims any, representation, endorsement, guarantee and/or warranty, express or implied, regarding this document. Any use of this document or the information contained within this document is at your own risk.")
        p = self.document.add_paragraph()
        p.add_run("Limitation of Liability").bold = True
        self.document.add_paragraph("Under no circumstances will any of the Dealwise Advisors Parties be liable for any loss or damage caused by your reliance on the information contained in this document. Because some jurisdictions do not allow the exclusion or limitation of liability for negligence, consequential, incidental or other types of damages, in such jurisdictions the Dealwise Parties’ liability is limited to the greatest extent permitted by law.")
        self.document.add_page_break()
        
        ### Preamble
        self.document.add_paragraph(f"""{datetime.now().strftime("%B %d, %Y")}
{loi.seller_name}
{loi.business_address}

Re: Proposed Transaction

Dear {loi.buyer_name}:
""")
        p = self.document.add_paragraph(f"The following sets forth the terms and conditions upon which an acquisition company to be established by {loi.legal_entity if loi.legal_entity else loi.buyer_name} and its affiliates (”")
        p.add_run("Buyer").bold = True
        p.add_run("”) would acquire ")
        p.add_run("all or substantially all of the assets of").bold = True 
        p.add_run(f"{loi.business_name}, a {loi.business_entity_type} (the “")
        p.add_run("Company").bold = True
        p.add_run("”). The acquisition of the Company and the related transactions are collectively referred to herein as the “")
        p.add_run("Transaction").bold = True
        p.add_run(".”")
        ### 1. Enterprise Valuation.
        p=self.document.add_paragraph(f"1.\t")
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.5), alignment=WD_TAB_ALIGNMENT.LEFT, leader=WD_TAB_LEADER.SPACES)
        r = p.add_run(f"Enterprise Valuation.")
        r.bold = True
        r.underline = True
        p.add_run(f"  Our proposal for the cash-free, debt-free enterprise valuation of the Company upon which the various aspects of the Transaction would be based is ${loi.purchase_price:,.0f} (the “")
        p.add_run("Company Enterprise Valuation").bold = True
        p.add_run(f"”). We have assumed in establishing the Company Enterprise Valuation that the Company will have run-rate revenue of approximately ${loi.biz_revenue:,.0f} and run-rate EBITDA of approximately ${loi.biz_ebitda:,.0f}, and that Buyer will receive a full step up in the tax basis of the assets in connection with the consummation of the Transaction.")

        ### 2. Purchase Price; Adjustment to PUrchase Price.
        note_amount = loi.purchase_price * loi.note_percent
        equity_rollover_amount = loi.purchase_price * loi.equity_rollover_percent
        closing_cash_amount = loi.purchase_price - note_amount - equity_rollover_amount
        next_num = "ii"
        p=self.document.add_paragraph(f"2.\t")
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.5), alignment=WD_TAB_ALIGNMENT.LEFT, leader=WD_TAB_LEADER.SPACES)
        r = p.add_run(f"Purchase Price; Adjustment to Purchase Price.")
        r.bold = True
        r.underline = True
        p.add_run(f"  At closing, Buyer will acquire the Company for a total aggregate purchase price equal to (a) the Company Enterprise Valuation adjusted as described herein, minus (b) the amount of any net indebtedness of the Company at closing, including loans, capital leases, letters of credit, any payments arising out of the sale of the Company and any unpaid transaction fees (the “")
        p.add_run(f"Purchase Price").bold = True
        p.add_run(f"”). The Purchase Price will be adjusted up or down to the extent that the Company’s net working capital at closing is greater than or less than a normalized working capital target to be mutually agreed upon in the definitive agreements. The Purchase Price will be comprised of (i) a closing cash payment in the amount of ${closing_cash_amount:,.0f} (adjusted as described herein)")

        if equity_rollover_amount > 0:
            p.add_run(f", ({next_num}) ${equity_rollover_amount:,.0f}, which will be subordinate to Buyer’s senior credit facility, be reinvested in Buyer on the same terms as capital contributed by Buyer’s other investors at the closing")
            next_num = "iii"
        
        if loi.note_percent > 0:
            p.add_run(f", ({next_num}) ${note_amount:,.0f}, which will bear interest at a rate of {loi.note_interest_rate * 100:.0f}% per year for {int(loi.note_term)} years")
            if loi.bullet_payment_anniversary:
                p.add_run(f" (interest will be payable quarterly in arrears) and be paid in a bullet payment on the {inflect_engine.ordinal(loi.bullet_payment_anniversary)} anniversary of the closing date")
            elif loi.note_standby > 0:
                p.add_run(f" and be repaid in equal monthly installments of principal and interest, following a {loi.note_standby}-year standby period where no payments of principal or interest are required")
            else:
                p.add_run(f" and be repaid in equal monthly installments of principal and interest commencing immediately after the closing date")
            next_num = "iv"
        
        if loi.earnout_description:
            p.add_run(f", ({next_num}) an earnout payable in cash upon the achievement of certain performance milestones by the Company:  {loi.earnout_description}")

        p.add_run(". Nothing in the Definitive Agreement will prohibit or limit Buyer’s ability to freely operate the business following the closing.")

        ### 3. Terms of Definitive Agreement; Conditions.
        p=self.document.add_paragraph(f"3.\t")
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.5), alignment=WD_TAB_ALIGNMENT.LEFT, leader=WD_TAB_LEADER.SPACES)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.1), alignment=WD_TAB_ALIGNMENT.LEFT, leader=WD_TAB_LEADER.SPACES)
        r = p.add_run(f"Terms of Definitive Agreement; Conditions.")
        r.bold = True
        r.underline = True
        p = self.document.add_paragraph(f"\ta.\tBuyer and its counsel will prepare a definitive purchase agreement (the “")
        p.add_run("Definitive Agreement").bold = True
        p.add_run("”) and related documentation. The Definitive Agreement will be mutually acceptable to the parties and include, among other things, customary representations and warranties relating to all aspects of the Company’s business (including without limitation assets, liabilities, contracts, operations,customers and suppliers, employees and compliance with laws), and indemnifications, escrows and covenants.")
        next_letter = "b"
        if loi.escrow_percent > 0:
            escrow_amount = loi.purchase_price * loi.escrow_percent
            p = self.document.add_paragraph(f"\t{next_letter}.\tThe seller and the Company will indemnify Buyer after the closing for claims relating to (i) breaches of representations, warranties and covenants and (ii) liabilities relating to the Company. At the closing, the parties will deposit ${escrow_amount:,.0f} ({loi.escrow_percent * 100:.0f})% of the Purchase Price into an escrow account as a source of security for such indemnification obligations. The non- fundamental representations and warranties shall survive until the eighteen- (18-) month anniversary of the closing, at which time all funds remaining in such escrow account (other than those reserved for then- pending claims) will be released to the Seller. Indemnification claims resulting from breaches of non-fundamental representations and warranties will be subject to a ${loi.escrow_cap:,.0f} cap and a ${loi.escrow_tipping_basket:,.0f} tipping basket. The parties will also deposit an agreed upon amount of the Purchase Price into a working capital escrow account to secure a customary post-closing true-up mechanism.")
            next_letter = "c"
        
        p = self.document.add_paragraph(f"\t{next_letter}.\tThe Definitive Agreement will also include closing conditions such as (i) satisfactory completion by Buyer of its due diligence investigation of the Company, (ii) a ")
        p.add_run("five (5) year").bold = True
        p.add_run(" non-competition agreement and customary non-solicitation and no-hire agreements executed by each of the sellers, (iii) receipt of all necessary approvals and consents and (iv) there having occurred no material adverse effect (which shall be defined in the Definitive Agreement) on the Company or its business since the date hereof.")

        ### 4. Key Due Diligence Items. 
        p=self.document.add_paragraph(f"4.\t")
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.5), alignment=WD_TAB_ALIGNMENT.LEFT, leader=WD_TAB_LEADER.SPACES)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.1), alignment=WD_TAB_ALIGNMENT.LEFT, leader=WD_TAB_LEADER.SPACES)
        r = p.add_run(f"Key Due Diligence Items.")
        r.bold = True
        r.underline = True
        p.add_run("  We have based this proposal upon our review to date, however, prior to becoming bound to the Definitive Agreement, we will need to complete the following additional matters of due diligence:")
        self.document.add_paragraph(f"\ta.\tManagement meetings and facility tours;")
        self.document.add_paragraph(f"\tb.\tTypical financial and legal due diligence{';' if len(loi.due_diligence_items) > 0 else '; and'}")
        self.document.add_paragraph(f"\tc.\tOther information required by financing sources{'; and' if len(loi.due_diligence_items) > 0 else '.'}")
        next_letter = "d"
        for i, item in enumerate(loi.due_diligence_items):
            self.document.add_paragraph(f"\t{next_letter}.\t{item}")
            next_letter = chr(ord(next_letter) + 1)

        ### 5. Due Diligence Access.
        p=self.document.add_paragraph(f"5.\t")
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.5), alignment=WD_TAB_ALIGNMENT.LEFT, leader=WD_TAB_LEADER.SPACES)
        r = p.add_run(f"Due Diligence Access.")
        r.bold = True
        r.underline = True
        p.add_run("  Promptly following the execution of this letter of intent, the Company will (and will cause its officers, directors, employees, agents and representatives to) provide Buyer and its representatives with full access during normal business hours and upon reasonable advance notice to the Company’s management, employees, customers, suppliers, financial statements, books and records, contracts, leases, operations, forecasts, tax records and other documents, other than information which is subject to attorney-client privilege in the manner and to the extent Buyer deems desirable in its sole discretion.")

        ### 6. Publicity.
        p=self.document.add_paragraph(f"6.\t")
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.5), alignment=WD_TAB_ALIGNMENT.LEFT, leader=WD_TAB_LEADER.SPACES)
        r = p.add_run(f"Publicity.")
        r.bold = True
        r.underline = True
        p.add_run("  Without the prior written consent of the other parties, neither Buyer, the Company, nor the owners of the Company will, and each will direct its representatives not to, make any public communication with respect to, or otherwise disclose or permit the disclosure of the existence of, discussions regarding a possible transaction between the parties hereto, the existence of this letter of intent, or any of the terms, conditions or other aspects of the transaction proposed in this letter of intent. Notwithstanding the foregoing, any party (and any employee, representative or other agent of any party) may disclose to any and all persons the tax treatment, tax structure and tax consequences of the transactions contemplated herein and all materials of any kind (including opinions or other tax analyses) that are provided to it relating to such tax treatment, tax structure and tax consequences.")

        ### 7. Exclusivity.
        p=self.document.add_paragraph(f"7.\t")
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.5), alignment=WD_TAB_ALIGNMENT.LEFT, leader=WD_TAB_LEADER.SPACES)
        r = p.add_run(f"Exclusivity.")
        r.bold = True
        r.underline = True
        p.add_run(f"  The Company agrees, in consideration of the substantial expenditure of time, effort and expense to be undertaken by Buyer in the investigation, preparation of documents and other activities related to the transactions contemplated by this letter of intent, that during the period commencing with the date of this letter of intent and ending {inflect_engine.number_to_words(loi.exclusivity_days)} ({loi.exclusivity_days}) days thereafter (the “")
        p.add_run("Exclusivity Period").bold = True
        p.add_run("”), neither the Company nor any of its stockholders/members, officers, directors/managers, employees, agents or representatives, will: (a) solicit, initiate or encourage submission of proposals or offers, or accept any proposals or offers, or enter into negotiations or discussions with any other person or persons (including, without limitation, any financial or other advisors) with regard to the merger or combination with, or sale or any other form of disposition of, the assets or stock/membership interests or other security interests of the Company or any part thereof (“")
        p.add_run("Acquisition Proposals").bold = True
        p.add_run("”), and that any such negotiations or discussions that may now be in process will be terminated; or (b) furnish to any other person any information with respect to, or otherwise cooperate in any way with or assist, facilitate or encourage, any Acquisition Proposal. The Company will promptly notify Buyer regarding any contact between the Company, its stockholders/equityholders or their respective representatives and any other person regarding any Acquisition Proposal. Unless either party provides the other written notice at least five (5) business days prior to the expiration of the Exclusivity Period, the Exclusivity Period shall automatically extend for successive thirty (30) day periods.")

        ### 8. Expenses
        p=self.document.add_paragraph(f"8.\t")
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.5), alignment=WD_TAB_ALIGNMENT.LEFT, leader=WD_TAB_LEADER.SPACES)
        r = p.add_run(f"Expenses.")
        r.bold = True
        r.underline = True
        p.add_run(f"  Buyer and the Company will each bear their respective expenses in negotiating and executing this letter of intent and the Definitive Agreement and completing the transactions contemplated thereby, including conducting due diligence and obtaining financing.")

        ### 9. Effect of Letter.
        p=self.document.add_paragraph(f"9.\t")
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.5), alignment=WD_TAB_ALIGNMENT.LEFT, leader=WD_TAB_LEADER.SPACES)
        r = p.add_run(f"Effect of Letter.")
        r.bold = True
        r.underline = True
        p.add_run(f"  It is understood that this letter of intent is intended to be a summary of terms on which the parties expect to proceed and is not intended to create and does not create any legally binding obligations, except as set forth in Paragraphs 6, 7, 8, 9, 10, 11 and 12 (the “")
        p.add_run(f"Binding Provisions").bold = True
        p.add_run(f"”). The parties intend that legally binding obligations with respect to the purchase and sale of the Company will arise only pursuant to the execution of a Definitive Agreement.")

        ### 10. Termination.
        p=self.document.add_paragraph(f"10.\t")
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.5), alignment=WD_TAB_ALIGNMENT.LEFT, leader=WD_TAB_LEADER.SPACES)
        r = p.add_run(f"Termination.")
        r.bold = True
        r.underline = True
        p.add_run(f"  Unless otherwise agreed in writing by all of the parties, Paragraph 7 will automatically terminate after the expiration of the Exclusivity Period; provided, however, that any termination shall not affect the liability of a party for breach of Paragraph 7 prior to the termination. Notwithstanding the foregoing, unless expressly terminated or modified in writing by all of the parties, Paragraphs 6, 8, 9, 10, 11 and 12 shall survive any termination of this letter of intent.")
        
        ### 11. Entire Agreement.
        p=self.document.add_paragraph(f"11.\t")
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.5), alignment=WD_TAB_ALIGNMENT.LEFT, leader=WD_TAB_LEADER.SPACES)
        r = p.add_run(f"Entire Agreement.")
        r.bold = True
        r.underline = True
        p.add_run(f"  The Binding Provisions constitute the entire agreement between the parties and supersede all prior oral or written agreements or understandings of the parties with regard to the subject matter hereof, except for any and all non-disclosure agreements previously entered into between the parties, which shall remain in full force and effect.")

        ### 12. Governing Law.
        p=self.document.add_paragraph(f"12.\t")
        p.paragraph_format.tab_stops.add_tab_stop(Inches(0.5), alignment=WD_TAB_ALIGNMENT.LEFT, leader=WD_TAB_LEADER.SPACES)
        r = p.add_run(f"Governing Law.")
        r.bold = True
        r.underline = True
        p.add_run(f"  The Binding Provisions will be governed by and construed under the laws of the State of {loi.governing_law} without regard to conflicts of laws principles.")

        ### Signature
        p = self.document.add_paragraph("[")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("Signature Page to Follow").italic = True
        p.add_run("]")
        self.document.add_page_break()
        p = self.document.add_paragraph(f"""If you are in agreement with the foregoing, please sign and return one copy of this letter of intent on behalf of the Company. We look forward to working with you.

Sincerely,
""")
        p.add_run(loi.legal_entity if loi.legal_entity else "").bold = True
        p.add_run(f"""


By: _______________________
Name: {loi.buyer_name}
Title: {loi.buyer_title}



ACCEPTED AND AGREED:
""")

        p.add_run({loi.business_name}).bold = True
        p.add_run(f"""


By: _______________________
Name: {loi.seller_name}
Title: _______________________
Date: _______________________
""")
    def save(self, path: str):
        self.document.save(path)