import requests
import json
from typing import Dict, List, Optional
from models.models import Source, AppConfig, Document, DocumentMetadata, DataConnector, AuthorizationResult
from appstatestore.statestore import StateStore
from io import BytesIO
from PyPDF2 import PdfReader
import docx
import uuid
import os

BASE_URL = "https://api.dropboxapi.com"
APP_KEY = os.environ.get("DROPBOX_APP_KEY")
APP_SECRET = os.environ.get("DROPBOX_APP_SECRET")


class DropboxConnector(DataConnector):
    source_type: Source = Source.dropbox
    connector_id: int = 7
    config: AppConfig
    folder_name: str

    def __init__(self, config: AppConfig, folder_name: str):
        super().__init__(config=config, folder_name=folder_name)

    def check_valid_access_token(self, access_token: str) -> bool:
        headers = {
            'Authorization': f'Bearer {access_token}',
            'Content-Type': 'application/json',
        }

        json_data = {
            'query': 'foo',
        }
        response = requests.post(f'{BASE_URL}/2/check/user', headers=headers, json=json_data)
        if response.status_code == 200:
            return True
        else:
            return False

    async def authorize(self, redirect_uri: str, auth_code: Optional[str]) -> str | None:
        if APP_KEY and APP_SECRET and auth_code:
            data = {
                'code': auth_code,
                'grant_type': 'authorization_code',
                'client_id': APP_KEY,
                'client_secret': APP_SECRET,
            }

            response = requests.post('https://api.dropbox.com/oauth2/token', data=data)
            if response.status_code == 200:
                response_data = response.json()
                refresh_token = response_data.get('refresh_token')

                creds = {
                    "refresh_token": refresh_token
                }

                creds_string = json.dumps(creds)
                StateStore().save_credentials(self.config, creds_string, self)
        else:
            if APP_KEY:
                auth_url = f"https://www.dropbox.com/oauth2/authorize?client_id={APP_KEY}&token_access_type=offline&response_type=code"
                return auth_url


    def get_new_access_token(self, app_key: str, app_secret: str, refresh_token: str):
        data = {
            'grant_type': 'refresh_token',
            'refresh_token': refresh_token,
            'client_id': app_key,
            'client_secret': app_secret,
        }

        response = requests.post('https://api.dropbox.com/oauth2/token', data=data)
        if response.status_code == 200:
            res_data = response.json()
            return res_data.get('access_token')
        else:
            return None

    def get_all_files_under_folder(self, access_token: str, folder_name: str):
        headers = {
            'Authorization': f"Bearer {access_token}",
            'Content-Type': 'application/json',
        }

        json_data = {
            'include_deleted': False,
            'include_has_explicit_shared_members': False,
            'include_media_info': False,
            'include_mounted_folders': True,
            'include_non_downloadable_files': True,
            'path': f"/{folder_name}",
            'recursive': True,
        }

        res = requests.post(f'{BASE_URL}/2/files/list_folder', headers=headers, json=json_data)
        data = res.json()
        files = data.get('entries')

        file_metadata = []
        for doc in files:
            if doc.get('.tag') == "file":
                document_name = doc.get('name')
                file_type = document_name.split('.')[-1]
                if doc.get('is_downloadable'):
                    document_path = doc.get('path_lower')
                    file_metadata.append((document_name, document_path, file_type))

        return file_metadata

    def extract_text_from_document(self, access_token: str, document_path: str, file_type: str):
        path = {"path": document_path}
        path = json.dumps(path)
        headers = {
            'Authorization': f'Bearer {access_token}',
            'Dropbox-API-Arg': path,
        }

        response = requests.post('https://content.dropboxapi.com/2/files/download', headers=headers)
        if response.status_code == 200:
            content = response.content
            if content:
                text = ""
                if file_type == "pdf":
                    bytes_data = BytesIO(content)
                    reader = PdfReader(bytes_data)
                    for num in range(len(reader.pages)):
                        page = reader.pages[num]
                        page_text = page.extract_text()
                        text += page_text

                elif file_type == "docx" or file_type == "doc":
                    bytes_data = BytesIO(content)
                    doc = docx.Document(bytes_data)
                    for p in doc.paragraphs:
                        text += p.text

                elif file_type == "txt":
                    text = content.decode("utf-8")
                return text
        return None



    async def load(self, source_id: str) -> List[Document]:
        credential_string = StateStore().load_credentials(self.config, self)
        credential_json = json.loads(credential_string)

        documents: List[Document] = []

        # app_key = credential_json.get('app_key')
        # app_secret = credential_json.get('app_secret')
        # refresh_token = credential_json.get('refresh_token')

        app_key = APP_KEY
        app_secret = APP_SECRET
        refresh_token = credential_json.get('refresh_token')

        access_token = self.get_new_access_token(app_key, app_secret, refresh_token)

        if not access_token:
            return documents

        files = self.get_all_files_under_folder(access_token, self.folder_name)

        for file_name, file_path, file_type in files:
            text = self.extract_text_from_document(access_token, file_path, file_type)
            if text:
                documents.append(
                    Document(
                        title=file_name,
                        text=text,
                        url=file_path,
                        source_type=Source.dropbox,
                        metadata=DocumentMetadata(
                            document_id=str(uuid.uuid4()),
                            source_id=source_id,
                            tenant_id=self.config.tenant_id
                        )
                    )
                )

        return documents
